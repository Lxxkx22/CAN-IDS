import re
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from collections import defaultdict

def parse_test_log_to_word(log_path, output_path):
    """
    解析测试日志文件并生成Word文档，按测试类分组显示测试结果
    """
    # 创建Word文档
    doc = Document()
    doc.add_heading('IDS4.0 测试结果报告', 0)
    
    # 读取日志文件
    with open(log_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 按测试类分组存储测试数据
    tests_by_class = defaultdict(list)
    
    # 使用正则表达式提取测试信息
    # 匹配模式：测试 #数字 -> 类名 -> 方法 -> 描述 -> 预期结果 -> 实际结果
    pattern = r'测试 #\d+\s*\n\s*类名: (.+?)\s*\n\s*方法: (.+?)\s*\n\s*描述: (.+?)\s*\n\s*预期结果: (.+?)\s*\n\s*实际结果: (.+?)\s*\n\s*状态:'
    
    matches = re.findall(pattern, content, re.DOTALL)
    
    for match in matches:
        class_name = match[0].strip()
        method_name = match[1].strip()
        description = match[2].strip()
        expected_result = match[3].strip()
        actual_result = match[4].strip()
        
        tests_by_class[class_name].append({
            'method': method_name,
            'description': description,
            'expected': expected_result,
            'actual': actual_result
        })
    
    # 为每个测试类创建表格
    for class_name, tests in sorted(tests_by_class.items()):
        # 添加测试类标题
        doc.add_heading(f'测试类名: {class_name}', level=1)
        
        # 创建表格 (4列：测试方法名、描述、预期结果、实际结果)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 禁用表格自动调整，确保列宽固定
        table.autofit = False
        
        # 设置固定列宽（使用厘米单位更精确）
        col_widths = [Cm(4), Cm(6), Cm(6), Cm(3)]  # 总宽度19cm
        
        # 设置表格总宽度
        table.width = Cm(19)
        
        # 为每列设置固定宽度
        for i, width in enumerate(col_widths):
            table.columns[i].width = width
            # 强制设置列宽属性
            for cell in table.column_cells(i):
                cell.width = width
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '测试方法名'
        hdr_cells[1].text = '描述'
        hdr_cells[2].text = '预期结果'
        hdr_cells[3].text = '实际结果'
        
        # 设置表头样式
        for i, cell in enumerate(hdr_cells):
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 强制设置单元格宽度
            cell.width = col_widths[i]
            # 设置单元格内边距
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_widths[i].pt * 20)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            
            for run in cell.paragraphs[0].runs:
                run.bold = True
        
        # 添加测试数据行
        for test in tests:
            row_cells = table.add_row().cells
            row_cells[0].text = test['method']
            row_cells[1].text = test['description']
            row_cells[2].text = test['expected']
            row_cells[3].text = test['actual']
            
            # 设置单元格对齐方式和固定列宽
            for i, cell in enumerate(row_cells):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # 强制设置单元格宽度
                cell.width = col_widths[i]
                # 设置单元格内边距和宽度属性
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(col_widths[i].pt * 20)))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)
        
        # 在表格后添加空行
        doc.add_paragraph()
    
    # 保存文档
    doc.save(output_path)
    print(f"Word文档已生成: {output_path}")
    print(f"共处理了 {len(tests_by_class)} 个测试类")
    total_tests = sum(len(tests) for tests in tests_by_class.values())
    print(f"总测试方法数: {total_tests}")

if __name__ == "__main__":
    log_file_path = 'c:/Users/22254/PycharmProjects/IDS4.0/tests/comprehensive_test_log_20250608_163847_readable.txt'
    output_word_path = 'c:/Users/22254/PycharmProjects/IDS4.0/tests/测试结果报告.docx'
    
    parse_test_log_to_word(log_file_path, output_word_path)